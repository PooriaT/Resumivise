from docx import Document
import json
import re

def read_docx(file):
    doc = Document(file)
    section = doc.sections[0]
    header = [hr.text.strip() for hr in section.header.paragraphs if hr.text.strip()] 
    text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    data = json.dumps({'text': '\n'.join(header + text)})
    return data
 
def write_docx(filename, data):
    doc = Document()
    section = doc.sections[0]
    header = section.header

    header_text = ""
    pattern = [r'.*name.*', r'.*phone.*', r'.*location.*', r'.*email.*', r'.*link']
    combined_pattern = re.compile(fr'{pattern[0]}|{pattern[1]}|{pattern[2]}|{pattern[3]}|{pattern[4]}')
    for heading, text in data.items():
        heading = heading.strip()
        if combined_pattern.match(heading):
            header_text += f"{text} |"
        else:
            doc.add_heading(heading, level=1)
            if isinstance(text, list):
                for item in text:
                    if isinstance(item, dict):
                        for key, value in item.items():
                            doc.add_paragraph(f"{key}: {value}")
                    else:
                        doc.add_paragraph(str(item))
            else:
                if isinstance(text, dict):
                    for key, value in text.items():
                        doc.add_paragraph(f"{key}: {value}")
                else:
                    doc.add_paragraph(str(text))
            doc.add_paragraph()

    header.paragraphs[0].text = header_text
    doc.save(filename)
