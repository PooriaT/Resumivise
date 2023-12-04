from docx import Document
import json

def read_docx(file):
    doc = Document(file)
    section = doc.sections[0]
    header = [hr.text.strip() for hr in section.header.paragraphs if hr.text.strip()] 
    text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    data = json.dumps({'text': '\n'.join(header + text)})
    return data
    
# NEED TO BE REVISED
def write_docx(file, data):
    doc = Document()

    doc.add_heading('JSON Data', level=1)
    doc.add_paragraph(json.dumps(data, indent=4))

    doc.save(file)